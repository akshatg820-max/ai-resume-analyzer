
import os
import re
import json
import hashlib
from io import BytesIO
from collections import Counter

import streamlit as st
from dotenv import load_dotenv

# ------------------------------------------------------------
# Optional dependencies
# ------------------------------------------------------------
try:
    from google import genai
except ImportError:
    genai = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None


# ------------------------------------------------------------
# Configuration
# ------------------------------------------------------------
load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "").strip()
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash").strip()

st.set_page_config(
    page_title="AI Resume Analyzer",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ------------------------------------------------------------
# Professional UI styling
# ------------------------------------------------------------
st.markdown(
    """
    <style>
        :root {
            --bg: #07111f;
            --panel: #0e1b2d;
            --panel-2: #101f34;
            --border: #263a52;
            --text: #f5f7fb;
            --muted: #9fb0c5;
            --blue: #38bdf8;
            --blue-2: #0ea5e9;
            --green: #22c55e;
            --yellow: #f59e0b;
            --red: #ef4444;
        }

        .stApp {
            background:
                radial-gradient(circle at 85% 8%, rgba(14,165,233,.10), transparent 25%),
                radial-gradient(circle at 20% 0%, rgba(59,130,246,.08), transparent 22%),
                var(--bg);
            color: var(--text);
        }

        [data-testid="stSidebar"] {
            background: #06101d;
            border-right: 1px solid #17283d;
        }

        [data-testid="stSidebar"] * {
            color: #e8edf5;
        }

        .block-container {
            max-width: 1450px;
            padding-top: 1.7rem;
            padding-bottom: 3rem;
        }

        h1, h2, h3, h4 {
            letter-spacing: -0.02em;
        }

        .hero {
            padding: 1.8rem 2rem;
            border: 1px solid rgba(56,189,248,.16);
            border-radius: 22px;
            background:
                linear-gradient(135deg, rgba(18,40,68,.94), rgba(10,26,45,.88));
            box-shadow: 0 18px 60px rgba(0,0,0,.22);
            margin-bottom: 1.5rem;
        }

        .hero-title {
            font-size: 2.2rem;
            font-weight: 800;
            margin: 0;
        }

        .hero-sub {
            color: #afbdd0;
            margin-top: .5rem;
            font-size: 1.02rem;
        }

        .badge {
            display: inline-block;
            padding: .35rem .75rem;
            border-radius: 999px;
            background: rgba(14,165,233,.10);
            border: 1px solid rgba(56,189,248,.25);
            color: #93ddff;
            font-size: .83rem;
            font-weight: 700;
            margin-top: .9rem;
        }

        .panel {
            background: linear-gradient(180deg, rgba(16,31,52,.96), rgba(11,25,42,.96));
            border: 1px solid var(--border);
            border-radius: 18px;
            padding: 1.15rem 1.2rem;
            box-shadow: 0 14px 40px rgba(0,0,0,.16);
        }

        .section-title {
            font-size: 1.45rem;
            font-weight: 800;
            margin: .1rem 0 .25rem 0;
        }

        .section-sub {
            color: var(--muted);
            font-size: .93rem;
            margin-bottom: .8rem;
        }

        .metric-card {
            background: linear-gradient(180deg, #11223a, #0d1c30);
            border: 1px solid #29415e;
            border-radius: 16px;
            padding: 1rem;
            min-height: 120px;
        }

        .metric-label {
            color: #9fb0c5;
            font-size: .85rem;
        }

        .metric-value {
            font-size: 2rem;
            font-weight: 800;
            margin-top: .25rem;
        }

        .metric-note {
            color: #b7c5d7;
            font-size: .78rem;
            margin-top: .15rem;
        }

        .skill-pill {
            display: inline-block;
            padding: .36rem .68rem;
            margin: .2rem .2rem .2rem 0;
            border-radius: 999px;
            border: 1px solid rgba(34,197,94,.30);
            background: rgba(34,197,94,.10);
            color: #8af0ae;
            font-size: .84rem;
            font-weight: 700;
        }

        .area-pill {
            display: inline-block;
            padding: .34rem .62rem;
            margin: .2rem .2rem .2rem 0;
            border-radius: 999px;
            border: 1px solid rgba(99,102,241,.35);
            background: rgba(99,102,241,.12);
            color: #c4c6ff;
            font-size: .80rem;
            font-weight: 700;
        }

        .status-good {
            display: inline-block;
            padding: .38rem .75rem;
            border-radius: 999px;
            color: #86efac;
            background: rgba(34,197,94,.10);
            border: 1px solid rgba(34,197,94,.28);
            font-weight: 700;
        }

        .status-mid {
            display: inline-block;
            padding: .38rem .75rem;
            border-radius: 999px;
            color: #fde68a;
            background: rgba(245,158,11,.10);
            border: 1px solid rgba(245,158,11,.28);
            font-weight: 700;
        }

        .status-low {
            display: inline-block;
            padding: .38rem .75rem;
            border-radius: 999px;
            color: #fca5a5;
            background: rgba(239,68,68,.10);
            border: 1px solid rgba(239,68,68,.28);
            font-weight: 700;
        }

        .match-box {
            background: linear-gradient(135deg, rgba(14,165,233,.12), rgba(59,130,246,.05));
            border: 1px solid rgba(56,189,248,.22);
            border-radius: 16px;
            padding: 1rem 1.1rem;
        }

        .match-score {
            font-size: 2.7rem;
            font-weight: 850;
            line-height: 1;
        }

        .tiny {
            color: #93a5ba;
            font-size: .82rem;
        }

        .footer {
            margin-top: 2rem;
            padding-top: 1rem;
            border-top: 1px solid #22354b;
            text-align: center;
            color: #74869b;
            font-size: .82rem;
        }

        .score-ring {
            width: 132px;
            height: 132px;
            border-radius: 50%;
            display: grid;
            place-items: center;
            position: relative;
            flex: 0 0 auto;
        }

        .score-ring::after {
            content: "";
            position: absolute;
            width: 102px;
            height: 102px;
            border-radius: 50%;
            background: #0d1c30;
            border: 1px solid #29415e;
        }

        .score-ring-inner {
            position: relative;
            z-index: 2;
            text-align: center;
        }

        .score-ring-number {
            font-size: 2rem;
            font-weight: 850;
            line-height: 1;
        }

        .score-ring-label {
            color: #9fb0c5;
            font-size: .72rem;
            margin-top: .25rem;
        }

        .feature-card {
            background: linear-gradient(180deg, rgba(16,31,52,.92), rgba(9,23,39,.94));
            border: 1px solid #263f5c;
            border-radius: 16px;
            padding: 1rem;
            transition: transform .18s ease, border-color .18s ease, box-shadow .18s ease;
        }

        .feature-card:hover {
            transform: translateY(-2px);
            border-color: rgba(56,189,248,.55);
            box-shadow: 0 12px 30px rgba(0,0,0,.18);
        }

        .match-stat {
            background: rgba(6,16,29,.45);
            border: 1px solid rgba(56,189,248,.14);
            border-radius: 12px;
            padding: .75rem;
            text-align: center;
        }

        .match-stat-value {
            font-size: 1.25rem;
            font-weight: 800;
        }

        .match-stat-label {
            color: #8fa3bb;
            font-size: .72rem;
            margin-top: .15rem;
        }

        .priority-high {
            border-left: 4px solid #ef4444;
            background: rgba(239,68,68,.07);
            padding: .75rem 1rem;
            border-radius: 0 12px 12px 0;
            margin-bottom: .65rem;
        }

        .priority-medium {
            border-left: 4px solid #f59e0b;
            background: rgba(245,158,11,.07);
            padding: .75rem 1rem;
            border-radius: 0 12px 12px 0;
            margin-bottom: .65rem;
        }

        .priority-low {
            border-left: 4px solid #22c55e;
            background: rgba(34,197,94,.07);
            padding: .75rem 1rem;
            border-radius: 0 12px 12px 0;
            margin-bottom: .65rem;
        }

        .empty-state {
            text-align: center;
            padding: 2.2rem 1rem;
            border: 1px dashed #29415e;
            border-radius: 16px;
            background: rgba(10,24,41,.5);
        }

        @media (max-width: 900px) {
            .block-container {
                padding-left: 1rem;
                padding-right: 1rem;
            }

            .hero {
                padding: 1.35rem;
            }

            .hero-title {
                font-size: 1.8rem;
            }
        }


        /* ============================================================
           PREMIUM ANIMATED THEME
           ============================================================ */

        :root {
            --bg: #070b18;
            --panel: #10172b;
            --panel-2: #151d36;
            --border: rgba(129, 140, 248, .22);
            --text: #f8fafc;
            --muted: #a8b3c7;
            --primary: #8b5cf6;
            --primary-2: #6366f1;
            --cyan: #22d3ee;
            --pink: #ec4899;
            --green: #34d399;
            --orange: #f59e0b;
        }

        .stApp {
            background:
                radial-gradient(circle at 10% 10%, rgba(139,92,246,.18), transparent 24%),
                radial-gradient(circle at 88% 12%, rgba(34,211,238,.13), transparent 25%),
                radial-gradient(circle at 55% 90%, rgba(236,72,153,.09), transparent 30%),
                #070b18 !important;
        }

        /* Floating ambient blobs */
        .stApp::before,
        .stApp::after {
            content: "";
            position: fixed;
            width: 280px;
            height: 280px;
            border-radius: 50%;
            pointer-events: none;
            z-index: 0;
            filter: blur(80px);
            opacity: .18;
            animation: floatBlob 10s ease-in-out infinite alternate;
        }

        .stApp::before {
            top: 12%;
            left: 20%;
            background: #8b5cf6;
        }

        .stApp::after {
            right: 8%;
            bottom: 8%;
            background: #22d3ee;
            animation-delay: -4s;
        }

        @keyframes floatBlob {
            0% { transform: translate3d(-20px, 10px, 0) scale(1); }
            50% { transform: translate3d(35px, -25px, 0) scale(1.12); }
            100% { transform: translate3d(-10px, 30px, 0) scale(.92); }
        }

        @keyframes fadeUp {
            from { opacity: 0; transform: translateY(22px); }
            to { opacity: 1; transform: translateY(0); }
        }

        @keyframes fadeIn {
            from { opacity: 0; }
            to { opacity: 1; }
        }

        @keyframes glowPulse {
            0%, 100% { box-shadow: 0 0 0 rgba(139,92,246,0); }
            50% { box-shadow: 0 0 32px rgba(139,92,246,.22); }
        }

        @keyframes gradientShift {
            0% { background-position: 0% 50%; }
            50% { background-position: 100% 50%; }
            100% { background-position: 0% 50%; }
        }

        @keyframes shine {
            0% { transform: translateX(-120%); }
            100% { transform: translateX(120%); }
        }

        @keyframes pop {
            0% { transform: scale(.94); opacity: 0; }
            70% { transform: scale(1.025); opacity: 1; }
            100% { transform: scale(1); opacity: 1; }
        }

        .hero {
            position: relative;
            overflow: hidden;
            background:
                linear-gradient(125deg,
                    rgba(99,102,241,.22),
                    rgba(139,92,246,.16),
                    rgba(34,211,238,.10),
                    rgba(236,72,153,.10));
            background-size: 300% 300%;
            animation: gradientShift 12s ease infinite, glowPulse 5s ease-in-out infinite;
            border: 1px solid rgba(139,92,246,.38) !important;
            box-shadow:
                0 25px 80px rgba(0,0,0,.30),
                inset 0 1px rgba(255,255,255,.06);
        }

        .hero::after {
            content: "";
            position: absolute;
            top: 0;
            left: 0;
            width: 35%;
            height: 100%;
            background: linear-gradient(
                90deg,
                transparent,
                rgba(255,255,255,.09),
                transparent
            );
            transform: translateX(-120%);
            animation: shine 7s ease-in-out infinite;
        }

        .hero-title,
        .hero-sub,
        .badge {
            position: relative;
            z-index: 2;
        }

        .hero-title {
            animation: fadeUp .7s ease both;
        }

        .hero-sub {
            animation: fadeUp .7s .12s ease both;
        }

        .badge {
            animation: pop .7s .25s ease both;
            background: linear-gradient(90deg, rgba(139,92,246,.18), rgba(34,211,238,.15)) !important;
            border-color: rgba(139,92,246,.45) !important;
            color: #ddd6fe !important;
        }

        .metric-card,
        .panel,
        .feature-card,
        .match-box,
        .empty-state {
            animation: fadeUp .65s ease both;
            backdrop-filter: blur(12px);
            -webkit-backdrop-filter: blur(12px);
        }

        .metric-card {
            background:
                linear-gradient(145deg, rgba(21,29,54,.95), rgba(10,16,32,.95)) !important;
            border-color: rgba(99,102,241,.28) !important;
            position: relative;
            overflow: hidden;
            transition: transform .25s ease, border-color .25s ease, box-shadow .25s ease;
        }

        .metric-card::before {
            content: "";
            position: absolute;
            inset: 0;
            background: linear-gradient(120deg, transparent, rgba(255,255,255,.05), transparent);
            transform: translateX(-120%);
            transition: transform .6s ease;
        }

        .metric-card:hover {
            transform: translateY(-7px) scale(1.015);
            border-color: rgba(34,211,238,.55) !important;
            box-shadow: 0 18px 45px rgba(0,0,0,.28), 0 0 28px rgba(99,102,241,.10);
        }

        .metric-card:hover::before {
            transform: translateX(120%);
        }

        .metric-value {
            background: linear-gradient(90deg, #fff, #c4b5fd, #67e8f9);
            background-size: 200% auto;
            -webkit-background-clip: text;
            background-clip: text;
            color: transparent;
            animation: gradientShift 5s linear infinite;
        }

        .section-title {
            color: #f8fafc;
        }

        .section-title::first-letter {
            filter: drop-shadow(0 0 8px rgba(139,92,246,.65));
        }

        .skill-pill,
        .area-pill {
            transition: transform .18s ease, box-shadow .18s ease, background .18s ease;
            animation: pop .45s ease both;
        }

        .skill-pill {
            background: linear-gradient(135deg, rgba(52,211,153,.13), rgba(34,211,238,.08)) !important;
            border-color: rgba(52,211,153,.36) !important;
            color: #a7f3d0 !important;
        }

        .area-pill {
            background: linear-gradient(135deg, rgba(139,92,246,.16), rgba(99,102,241,.10)) !important;
            border-color: rgba(129,140,248,.42) !important;
            color: #ddd6fe !important;
        }

        .skill-pill:hover,
        .area-pill:hover {
            transform: translateY(-3px) scale(1.04);
            box-shadow: 0 8px 24px rgba(99,102,241,.18);
        }

        .status-good {
            background: rgba(52,211,153,.11) !important;
            border-color: rgba(52,211,153,.38) !important;
            color: #6ee7b7 !important;
            box-shadow: 0 0 22px rgba(52,211,153,.10);
        }

        .status-mid {
            background: rgba(245,158,11,.11) !important;
            border-color: rgba(245,158,11,.38) !important;
            color: #fcd34d !important;
        }

        .status-low {
            background: rgba(244,63,94,.11) !important;
            border-color: rgba(244,63,94,.38) !important;
            color: #fda4af !important;
        }

        .match-box {
            background:
                linear-gradient(135deg, rgba(99,102,241,.14), rgba(34,211,238,.07), rgba(236,72,153,.06)) !important;
            border-color: rgba(34,211,238,.30) !important;
        }

        .match-score {
            background: linear-gradient(90deg, #67e8f9, #a78bfa, #f0abfc);
            background-size: 200% auto;
            -webkit-background-clip: text;
            background-clip: text;
            color: transparent;
            animation: gradientShift 4s linear infinite;
        }

        .match-stat {
            transition: transform .22s ease, border-color .22s ease;
        }

        .match-stat:hover {
            transform: translateY(-4px);
            border-color: rgba(34,211,238,.35);
        }

        .score-ring {
            animation: glowPulse 3.5s ease-in-out infinite, pop .7s ease both;
            filter: drop-shadow(0 0 18px rgba(99,102,241,.18));
        }

        /* Modern Streamlit buttons */
        .stButton > button {
            position: relative;
            overflow: hidden;
            background:
                linear-gradient(100deg, #4f46e5, #7c3aed, #0891b2, #4f46e5) !important;
            background-size: 300% 100% !important;
            border: 1px solid rgba(255,255,255,.13) !important;
            box-shadow: 0 8px 25px rgba(79,70,229,.20);
            transition: transform .2s ease, box-shadow .2s ease;
            animation: gradientShift 7s ease infinite;
        }

        .stButton > button:hover {
            transform: translateY(-3px);
            box-shadow: 0 14px 34px rgba(79,70,229,.35);
        }

        .stButton > button:active {
            transform: scale(.98);
        }

        .stButton > button::after {
            content: "";
            position: absolute;
            top: -50%;
            left: -80%;
            width: 35%;
            height: 200%;
            transform: rotate(20deg);
            background: rgba(255,255,255,.16);
            animation: shine 4.5s ease-in-out infinite;
        }

        /* Text inputs */
        .stTextArea textarea,
        .stTextInput input {
            transition: border-color .25s ease, box-shadow .25s ease, transform .2s ease;
            background: rgba(10,17,34,.88) !important;
            border-color: rgba(129,140,248,.27) !important;
        }

        .stTextArea textarea:focus,
        .stTextInput input:focus {
            border-color: #8b5cf6 !important;
            box-shadow: 0 0 0 2px rgba(139,92,246,.14), 0 0 25px rgba(99,102,241,.12) !important;
            transform: translateY(-1px);
        }

        /* Upload area */
        [data-testid="stFileUploaderDropzone"] {
            background:
                linear-gradient(135deg, rgba(99,102,241,.10), rgba(34,211,238,.06)) !important;
            border: 1px dashed rgba(129,140,248,.45) !important;
            transition: border-color .25s ease, background .25s ease, transform .25s ease;
        }

        [data-testid="stFileUploaderDropzone"]:hover {
            border-color: #22d3ee !important;
            background:
                linear-gradient(135deg, rgba(139,92,246,.16), rgba(34,211,238,.10)) !important;
            transform: translateY(-2px);
        }

        /* Progress bars */
        [data-testid="stProgress"] > div > div > div > div {
            background: linear-gradient(90deg, #6366f1, #8b5cf6, #22d3ee) !important;
            background-size: 200% 100% !important;
            animation: gradientShift 3s linear infinite;
        }

        /* Success/info messages */
        [data-testid="stAlert"] {
            animation: fadeUp .5s ease both;
            border-radius: 14px !important;
        }

        /* Expanders */
        div[data-testid="stExpander"] {
            transition: border-color .25s ease, transform .25s ease;
        }

        div[data-testid="stExpander"]:hover {
            border-color: rgba(129,140,248,.42) !important;
            transform: translateY(-1px);
        }

        /* Sidebar */
        [data-testid="stSidebar"] {
            background:
                radial-gradient(circle at 50% 5%, rgba(99,102,241,.12), transparent 30%),
                #050916 !important;
        }

        [data-testid="stSidebar"] .feature-card {
            animation: fadeIn .8s ease both;
        }

        /* Checkboxes */
        [data-testid="stCheckbox"] {
            transition: transform .18s ease;
        }

        [data-testid="stCheckbox"]:hover {
            transform: translateX(4px);
        }

        /* Reduce motion when user requests it at OS/browser level */
        @media (prefers-reduced-motion: reduce) {
            *,
            *::before,
            *::after {
                animation-duration: .01ms !important;
                animation-iteration-count: 1 !important;
                transition-duration: .01ms !important;
            }
        }

        div[data-testid="stExpander"] {
            border: 1px solid #22354b !important;
            border-radius: 14px !important;
            background: rgba(12,26,44,.7) !important;
        }

        .stButton > button {
            border-radius: 12px;
            border: 1px solid #2c4767;
            background: linear-gradient(180deg, #132b49, #0f2036);
            color: #f4f8fd;
            font-weight: 700;
        }

        .stButton > button:hover {
            border-color: #38bdf8;
            color: white;
        }

        .stTextArea textarea {
            background: #0c1727 !important;
            color: #edf4fc !important;
            border: 1px solid #2a415d !important;
            border-radius: 14px !important;
        }

        .stFileUploader {
            background: rgba(11,25,42,.5);
            border-radius: 14px;
        }

/* ============================================================
   V2 PREMIUM MOTION SYSTEM
   Glassmorphism + Aurora + Scanline + Microinteractions
   ============================================================ */

.block-container {
    position: relative;
    z-index: 1;
}

.stApp {
    background:
        radial-gradient(900px 450px at 12% -5%, rgba(124,58,237,.20), transparent 60%),
        radial-gradient(800px 430px at 100% 0%, rgba(6,182,212,.16), transparent 58%),
        radial-gradient(700px 430px at 70% 100%, rgba(236,72,153,.10), transparent 60%),
        #050816 !important;
    overflow-x: hidden;
}

.aurora-strip {
    height: 5px;
    border-radius: 999px;
    background: linear-gradient(90deg,#22d3ee,#6366f1,#a855f7,#ec4899,#22d3ee);
    background-size: 300% 100%;
    animation: auroraFlow 7s linear infinite;
    margin: -.2rem 0 1.1rem;
    box-shadow: 0 0 15px rgba(99,102,241,.28), 0 0 35px rgba(34,211,238,.14);
}

@keyframes auroraFlow {
    0% { background-position: 0% 50%; }
    100% { background-position: 300% 50%; }
}

@keyframes drift {
    0% { transform: translate3d(0,0,0) rotate(0deg); }
    33% { transform: translate3d(12px,-10px,0) rotate(2deg); }
    66% { transform: translate3d(-8px,8px,0) rotate(-2deg); }
    100% { transform: translate3d(0,0,0) rotate(0deg); }
}

@keyframes heartbeat {
    0%,100% { transform: scale(1); }
    50% { transform: scale(1.035); }
}

@keyframes scan {
    0% { transform: translateY(-115%); opacity: 0; }
    15% { opacity: .8; }
    85% { opacity: .25; }
    100% { transform: translateY(115%); opacity: 0; }
}

@keyframes textGlow {
    0%,100% { text-shadow: 0 0 0 rgba(167,139,250,0); }
    50% { text-shadow: 0 0 24px rgba(167,139,250,.22); }
}

@keyframes cardIn {
    0% { opacity: 0; transform: translateY(18px) scale(.985); }
    100% { opacity: 1; transform: translateY(0) scale(1); }
}

.hero { isolation: isolate; }

.hero::before {
    content: "";
    position: absolute;
    inset: 0;
    border-radius: inherit;
    background:
        radial-gradient(circle at 18% 30%, rgba(255,255,255,.11), transparent 22%),
        radial-gradient(circle at 82% 65%, rgba(34,211,238,.08), transparent 20%);
    pointer-events: none;
    z-index: 0;
    animation: drift 11s ease-in-out infinite;
}

.hero-title {
    animation: fadeUp .65s ease both, textGlow 5s ease-in-out 1.2s infinite;
}

.hero-sub {
    animation: fadeUp .65s .10s ease both;
}

.badge {
    animation: pop .65s .18s ease both, heartbeat 5s ease-in-out 1.3s infinite;
}

.workflow {
    display: grid;
    grid-template-columns: repeat(4, minmax(0,1fr));
    gap: .65rem;
    margin: .1rem 0 1.1rem;
}

.workflow-step {
    position: relative;
    overflow: hidden;
    padding: .75rem;
    border-radius: 14px;
    border: 1px solid rgba(129,140,248,.20);
    background: rgba(15,23,42,.55);
    backdrop-filter: blur(10px);
    transition: transform .22s ease, border-color .22s ease, box-shadow .22s ease;
    animation: cardIn .6s ease both;
}

.workflow-step:nth-child(2) { animation-delay: .08s; }
.workflow-step:nth-child(3) { animation-delay: .16s; }
.workflow-step:nth-child(4) { animation-delay: .24s; }

.workflow-step::after {
    content: "";
    position: absolute;
    inset: -80% 30% auto -20%;
    height: 120%;
    background: linear-gradient(115deg, transparent, rgba(255,255,255,.10), transparent);
    transform: rotate(12deg);
    animation: scan 4.8s ease-in-out infinite;
    pointer-events: none;
}

.workflow-step:hover {
    transform: translateY(-4px);
    border-color: rgba(34,211,238,.42);
    box-shadow: 0 14px 35px rgba(0,0,0,.22), 0 0 24px rgba(99,102,241,.10);
}

.workflow-num {
    display: inline-grid;
    place-items: center;
    width: 28px;
    height: 28px;
    border-radius: 50%;
    background: linear-gradient(135deg,#6366f1,#22d3ee);
    color: white;
    font-size: .75rem;
    font-weight: 900;
    box-shadow: 0 0 18px rgba(99,102,241,.22);
}

.workflow-label {
    margin-top: .45rem;
    font-size: .82rem;
    font-weight: 800;
}

.bento-card {
    position: relative;
    overflow: hidden;
    min-height: 155px;
    border-radius: 20px;
    border: 1px solid rgba(129,140,248,.22);
    background: linear-gradient(145deg, rgba(19,28,53,.84), rgba(7,12,25,.78));
    backdrop-filter: blur(14px);
    box-shadow: 0 16px 50px rgba(0,0,0,.18), inset 0 1px rgba(255,255,255,.05);
    padding: 1rem 1.1rem;
    animation: cardIn .55s ease both;
    transition: transform .22s ease, border-color .22s ease, box-shadow .22s ease;
}

.bento-card::before {
    content: "";
    position: absolute;
    width: 180px;
    height: 180px;
    right: -60px;
    top: -80px;
    border-radius: 50%;
    background: radial-gradient(circle, rgba(99,102,241,.24), transparent 65%);
    animation: drift 8s ease-in-out infinite;
}

.bento-card::after {
    content: "";
    position: absolute;
    inset: 0;
    border-radius: inherit;
    background: linear-gradient(120deg,transparent 0%,transparent 42%,rgba(255,255,255,.07) 50%,transparent 58%,transparent 100%);
    transform: translateX(-100%);
    animation: shine 8s ease-in-out infinite;
    pointer-events: none;
}

.bento-card:hover {
    transform: translateY(-5px);
    border-color: rgba(34,211,238,.34);
    box-shadow: 0 22px 55px rgba(0,0,0,.25), 0 0 28px rgba(99,102,241,.11);
}

.bento-label {
    position: relative;
    z-index: 2;
    color: #9fb0c5;
    font-size: .78rem;
    text-transform: uppercase;
    letter-spacing: .08em;
    font-weight: 850;
}

.bento-value {
    position: relative;
    z-index: 2;
    font-size: 2rem;
    line-height: 1;
    margin-top: .5rem;
    font-weight: 900;
    background: linear-gradient(90deg,#fff,#c4b5fd,#67e8f9);
    background-size: 200% auto;
    -webkit-background-clip: text;
    background-clip: text;
    color: transparent;
    animation: gradientShift 4s linear infinite;
}

.donut-wrap {
    display: flex;
    align-items: center;
    justify-content: center;
    min-height: 235px;
}

.donut {
    --pct: 88;
    width: 154px;
    height: 154px;
    border-radius: 50%;
    background: conic-gradient(from 220deg,#22d3ee 0%,#6366f1 calc(var(--pct) * 1%),#1d2940 calc(var(--pct) * 1%),#1d2940 100%);
    display: grid;
    place-items: center;
    position: relative;
    animation: heartbeat 4.8s ease-in-out infinite, cardIn .7s ease both;
    box-shadow: 0 0 30px rgba(99,102,241,.15), inset 0 0 20px rgba(0,0,0,.18);
}

.donut::before {
    content: "";
    position: absolute;
    width: 112px;
    height: 112px;
    border-radius: 50%;
    background: #0a1020;
    border: 1px solid rgba(129,140,248,.18);
}

.donut-content {
    position: relative;
    z-index: 2;
    text-align: center;
}

.donut-number {
    font-size: 2rem;
    font-weight: 900;
    color: #fff;
}

.donut-caption {
    color: #93a5ba;
    font-size: .75rem;
    margin-top: .1rem;
}

.radar-card {
    position: relative;
    min-height: 235px;
    padding: 1rem;
    border-radius: 20px;
    background: linear-gradient(145deg, rgba(15,23,42,.86), rgba(7,12,24,.84));
    border: 1px solid rgba(129,140,248,.22);
    overflow: hidden;
}

.radar-grid {
    display: grid;
    gap: .75rem;
    margin-top: .75rem;
}

.radar-row {
    display: grid;
    grid-template-columns: 130px 1fr 46px;
    gap: .7rem;
    align-items: center;
}

.radar-name {
    color: #dbe4f2;
    font-size: .79rem;
    font-weight: 750;
}

.radar-track {
    height: 9px;
    border-radius: 999px;
    background: rgba(51,65,85,.55);
    overflow: hidden;
}

.radar-fill {
    height: 100%;
    width: 0;
    border-radius: inherit;
    background: linear-gradient(90deg,#22d3ee,#6366f1,#a855f7);
    background-size: 200% 100%;
    animation: barGrow 1.35s cubic-bezier(.2,.8,.2,1) forwards, gradientShift 4s linear infinite;
    box-shadow: 0 0 12px rgba(99,102,241,.16);
}

@keyframes barGrow {
    from { width: 0; }
    to { width: var(--fill); }
}

.radar-score {
    color: #a5b4fc;
    font-size: .76rem;
    font-weight: 850;
    text-align: right;
}

.pulse-orb {
    width: 12px;
    height: 12px;
    display: inline-block;
    border-radius: 50%;
    background: #22d3ee;
    box-shadow: 0 0 0 0 rgba(34,211,238,.55);
    animation: pulseRing 2s infinite;
    margin-right: .35rem;
}

@keyframes pulseRing {
    0% { box-shadow: 0 0 0 0 rgba(34,211,238,.5); }
    70% { box-shadow: 0 0 0 13px rgba(34,211,238,0); }
    100% { box-shadow: 0 0 0 0 rgba(34,211,238,0); }
}

.floating-hint {
    border: 1px solid rgba(129,140,248,.20);
    background: rgba(15,23,42,.48);
    border-radius: 14px;
    padding: .7rem .85rem;
    font-size: .78rem;
    color: #a8b3c7;
    animation: drift 7s ease-in-out infinite;
}

.quick-action {
    position: relative;
    overflow: hidden;
    border-radius: 16px;
    border: 1px solid rgba(129,140,248,.24);
    background: linear-gradient(135deg, rgba(79,70,229,.16), rgba(6,182,212,.08));
    padding: .9rem 1rem;
    transition: transform .22s ease, box-shadow .22s ease, border-color .22s ease;
}

.quick-action:hover {
    transform: translateY(-4px) scale(1.01);
    border-color: rgba(34,211,238,.42);
    box-shadow: 0 16px 36px rgba(0,0,0,.22);
}

.quick-action::after {
    content: "";
    position: absolute;
    inset: -20% -40%;
    background: linear-gradient(110deg,transparent 40%,rgba(255,255,255,.10),transparent 60%);
    transform: translateX(-60%);
    animation: shine 6s ease-in-out infinite;
    pointer-events: none;
}

.priority-high,
.priority-medium,
.priority-low {
    transition: transform .2s ease, box-shadow .2s ease;
}

.priority-high:hover,
.priority-medium:hover,
.priority-low:hover {
    transform: translateX(5px);
    box-shadow: 0 9px 26px rgba(0,0,0,.16);
}

[data-testid="stFileUploaderDropzone"]::before {
    content: "DROP • SCAN • ANALYZE";
    display: block;
    font-size: .66rem;
    letter-spacing: .16em;
    font-weight: 900;
    color: #67e8f9;
    margin-bottom: .25rem;
    opacity: .8;
}

@media (max-width: 850px) {
    .workflow { grid-template-columns: repeat(2,minmax(0,1fr)); }
    .radar-row { grid-template-columns: 100px 1fr 40px; }
}

@media (prefers-reduced-motion: reduce) {
    *,*::before,*::after {
        animation-duration: .01ms !important;
        animation-iteration-count: 1 !important;
        scroll-behavior: auto !important;
        transition-duration: .01ms !important;
    }
}

    </style>
    """,
    unsafe_allow_html=True,
)

# ------------------------------------------------------------
# Skill database - career agnostic
# ------------------------------------------------------------
SKILLS_DATABASE = {
    # Programming / development
    "Python": ["python", "py"],
    "C": [r"\bc\b", "c programming", "c language"],
    "C++": ["c++", "cpp"],
    "Java": ["java"],
    "JavaScript": ["javascript", "js"],
    "HTML": ["html", "html5"],
    "CSS": ["css", "css3"],
    "React": ["react", "reactjs"],
    "Node.js": ["node.js", "nodejs"],
    "SQL": ["sql"],
    "Git": ["git", "version control"],
    "GitHub": ["github"],

    # Data / AI / ML
    "Machine Learning": ["machine learning", "machine-learning", "ml"],
    "Deep Learning": ["deep learning", "deep-learning"],
    "Artificial Intelligence": ["artificial intelligence", "ai"],
    "Data Analysis": ["data analysis", "data analytics", "data analyst"],
    "Data Science": ["data science", "data scientist"],
    "Data Structures": ["data structures", "data structure", "dsa"],
    "Algorithms": ["algorithms", "algorithm"],
    "Pandas": ["pandas"],
    "NumPy": ["numpy", "np"],
    "Scikit-learn": ["scikit-learn", "sklearn"],
    "TensorFlow": ["tensorflow"],
    "PyTorch": ["pytorch"],
    "Matplotlib": ["matplotlib"],
    "Seaborn": ["seaborn"],
    "Power BI": ["power bi", "powerbi"],
    "Tableau": ["tableau"],
    "Excel": ["excel", "microsoft excel", "ms excel"],

    # Engineering / cloud / tools
    "Docker": ["docker", "containerization"],
    "Kubernetes": ["kubernetes", "k8s"],
    "AWS": ["aws", "amazon web services"],
    "Azure": ["azure", "microsoft azure"],
    "Google Cloud": ["google cloud", "gcp"],
    "Linux": ["linux", "ubuntu"],
    "REST API": ["rest api", "restful api", "api development"],
    "FastAPI": ["fastapi"],
    "Flask": ["flask"],
    "Django": ["django"],

    # Business / operations
    "Project Management": ["project management", "project manager"],
    "Operations Management": ["operations management", "operations"],
    "Communication": ["communication", "verbal communication", "written communication"],
    "Leadership": ["leadership", "team leadership"],
    "Teamwork": ["teamwork", "team player", "team collaboration"],
    "Problem Solving": ["problem solving", "problem-solving"],
    "Critical Thinking": ["critical thinking"],
    "Time Management": ["time management"],
    "Customer Service": ["customer service", "customer support"],
    "Sales": ["sales", "sales management"],
    "Marketing": ["marketing"],
    "Digital Marketing": ["digital marketing"],
    "SEO": ["seo", "search engine optimization"],
    "Human Resources": ["human resources", "human resource", "hr"],
    "Recruitment": ["recruitment", "talent acquisition"],
    "Public Speaking": ["public speaking", "presentation skills"],
    "Research": ["research", "research skills"],

    # Logistics / supply chain
    "Logistics": ["logistics"],
    "Supply Chain Management": ["supply chain", "supply chain management"],
    "Inventory Management": ["inventory management", "inventory control", "inventory systems"],
    "Warehouse Operations": ["warehouse operations", "warehouse management", "warehousing"],
    "Picking": ["picking", "order picking"],
    "Packing": ["packing", "packaging"],
    "Shipping": ["shipping", "shipment"],
    "Receiving": ["receiving", "goods receiving"],
    "Distribution": ["distribution", "distribution center", "dc operations"],
    "Quality Control": ["quality control", "quality assurance", "qc"],
    "Lean Methodology": ["lean methodology", "lean manufacturing", "lean"],
    "Kaizen": ["kaizen"],
    "Gemba": ["gemba"],
    "5S": ["5s"],
    "Kanban": ["kanban"],
    "WMS": ["wms", "warehouse management system", "warehouse management systems"],
    "RF Scanners": ["rf scanner", "rf scanners", "radio frequency scanner"],
    "Forklift Operation": ["forklift", "forklift operation", "forklift certification"],
    "Pallet Jack": ["pallet jack", "electric pallet jack"],
    "Cycle Counting": ["cycle counting", "cycle count"],
    "OSHA": ["osha", "osha compliance", "osha standards"],
    "Order Fulfillment": ["order fulfillment", "fulfillment"],
    "Procurement": ["procurement", "purchasing"],
}

SKILL_AREAS = {
    "Programming & Development": {
        "Python", "C", "C++", "Java", "JavaScript", "HTML", "CSS",
        "React", "Node.js", "SQL", "Git", "GitHub", "Docker",
        "Kubernetes", "REST API", "FastAPI", "Flask", "Django"
    },
    "AI / ML & Data": {
        "Machine Learning", "Deep Learning", "Artificial Intelligence",
        "Data Analysis", "Data Science", "Data Structures", "Algorithms",
        "Pandas", "NumPy", "Scikit-learn", "TensorFlow", "PyTorch",
        "Matplotlib", "Seaborn", "Power BI", "Tableau"
    },
    "Business & Professional": {
        "Project Management", "Operations Management", "Communication",
        "Leadership", "Teamwork", "Problem Solving", "Critical Thinking",
        "Time Management", "Customer Service", "Sales", "Marketing",
        "Digital Marketing", "SEO", "Human Resources", "Recruitment",
        "Public Speaking", "Research"
    },
    "Logistics & Operations": {
        "Logistics", "Supply Chain Management", "Inventory Management",
        "Warehouse Operations", "Picking", "Packing", "Shipping",
        "Receiving", "Distribution", "Quality Control", "Lean Methodology",
        "Kaizen", "Gemba", "5S", "Kanban", "WMS", "RF Scanners",
        "Forklift Operation", "Pallet Jack", "Cycle Counting", "OSHA",
        "Order Fulfillment", "Procurement", "Excel"
    },
    "Cloud & Infrastructure": {
        "AWS", "Azure", "Google Cloud", "Linux"
    },
}

# Common skill suggestions are broad and intentionally career-agnostic.
COMMON_SUGGESTIONS = [
    "Problem Solving",
    "Communication",
    "Leadership",
    "Teamwork",
    "Project Management",
    "Critical Thinking",
    "Time Management",
    "Excel",
    "SQL",
    "Git",
    "Research",
    "Presentation",
]


# ------------------------------------------------------------
# Utility functions
# ------------------------------------------------------------
def normalize_text(text: str) -> str:
    text = text or ""
    text = text.replace("\x00", " ")
    return re.sub(r"\s+", " ", text).strip()


def safe_lower(text: str) -> str:
    return normalize_text(text).lower()


def keyword_in_text(text_lower: str, keyword: str) -> bool:
    """
    Handles common word/phrase matching without accidentally matching
    a single character like 'c' inside unrelated words.
    """
    keyword = keyword.lower().strip()
    if keyword == r"\bc\b":
        return bool(re.search(r"\bc\b", text_lower))
    if len(keyword) <= 2 and " " not in keyword:
        return bool(re.search(rf"\b{re.escape(keyword)}\b", text_lower))
    return keyword in text_lower


def detect_skills(text: str):
    text_lower = safe_lower(text)
    detected = []

    for skill, keywords in SKILLS_DATABASE.items():

        # Special handling for C to avoid false positives
        if skill == "C":
            c_patterns = [
                r"\bc\s+programming\b",
                r"\bc\s+language\b",
                r"\bprogramming\s+in\s+c\b",
                r"\bc\s+developer\b",
            ]

            if any(re.search(pattern, text_lower) for pattern in c_patterns):
                detected.append(skill)

            continue

        for keyword in keywords:
            if keyword_in_text(text_lower, keyword):
                detected.append(skill)
                break

    return detected


def detect_skill_areas(detected_skills):
    areas = []
    for area, area_skills in SKILL_AREAS.items():
        matched = [skill for skill in detected_skills if skill in area_skills]
        if matched:
            areas.append((area, len(matched)))
    return sorted(areas, key=lambda x: x[1], reverse=True)


def extract_pdf_text(uploaded_file):
    if PdfReader is None:
        return ""

    try:
        data = uploaded_file.getvalue()
        reader = PdfReader(BytesIO(data))
        pages = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(page_text.strip())
        return "\n".join(pages).strip()
    except Exception:
        return ""


def extract_docx_text(uploaded_file):
    if Document is None:
        return ""

    try:
        data = uploaded_file.getvalue()
        document = Document(BytesIO(data))
        parts = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts).strip()
    except Exception:
        return ""


def extract_text(uploaded_file):
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return extract_pdf_text(uploaded_file)
    if name.endswith(".docx"):
        return extract_docx_text(uploaded_file)
    return ""


def get_statistics(text: str):
    words = re.findall(r"\b[\w+#.-]+\b", text, flags=re.UNICODE)
    word_count = len(words)
    character_count = len(text)
    line_count = len([line for line in text.splitlines() if line.strip()])
    return word_count, character_count, line_count


def detect_sections(text: str):
    text_lower = safe_lower(text)

    section_keywords = {
        "Contact Information": ["email", "phone", "linkedin", "contact", "github"],
        "Summary / Profile": ["summary", "professional summary", "profile", "objective"],
        "Education": ["education", "academic", "degree", "bachelor", "master", "bca", "b.tech", "college", "university"],
        "Experience": ["experience", "employment", "work experience", "professional experience", "internship"],
        "Projects": ["projects", "project", "portfolio"],
        "Skills": ["skills", "technical skills", "technologies", "competencies"],
        "Certifications": ["certification", "certifications", "certificate", "licenses", "license"],
        "Achievements": ["achievements", "accomplishments", "awards"],
        "Languages": ["languages", "language proficiency"],
    }

    found = []
    for section, keywords in section_keywords.items():
        if any(keyword in text_lower for keyword in keywords):
            found.append(section)

    return found


def calculate_ats_score(text, detected_skills, sections):
    text_lower = safe_lower(text)
    score = 0

    # Skills: 25
    if len(detected_skills) >= 8:
        score += 25
    elif len(detected_skills) >= 5:
        score += 21
    elif len(detected_skills) >= 3:
        score += 16
    elif len(detected_skills) >= 1:
        score += 9

    # Section structure: 20
    core_sections = {"Contact Information", "Education", "Experience", "Skills"}
    core_found = len(core_sections.intersection(set(sections)))
    score += min(core_found * 5, 20)

    # Contact details: 15
    if re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text):
        score += 8
    if re.search(r"(?:\+?\d[\d\s().-]{7,}\d)", text):
        score += 7

    # Links: 10
    if "linkedin" in text_lower:
        score += 5
    if "github" in text_lower or "portfolio" in text_lower:
        score += 5

    # Projects / impact: 15
    if "project" in text_lower:
        score += 7
    if re.search(r"\b\d+(?:\.\d+)?\s?%", text):
        score += 4
    if re.search(r"\b(?:increased|decreased|reduced|improved|saved|grew|managed|delivered)\b", text_lower):
        score += 4

    # Readability / length: 15
    words = len(re.findall(r"\b[\w+#.-]+\b", text))
    if 180 <= words <= 700:
        score += 15
    elif 120 <= words < 180 or 700 < words <= 950:
        score += 10
    else:
        score += 5

    return min(score, 100)


def ats_label(score):
    if score >= 85:
        return "Excellent foundation", "good"
    if score >= 70:
        return "Good foundation", "good"
    if score >= 55:
        return "Fair — room for improvement", "mid"
    return "Needs improvement", "low"


def tokenize_keywords(text: str):
    stopwords = {
        "the", "and", "for", "with", "that", "this", "from", "your", "you",
        "are", "our", "will", "have", "has", "had", "into", "their", "they",
        "them", "then", "than", "been", "being", "who", "what", "when", "where",
        "how", "why", "can", "may", "must", "not", "all", "any", "use", "used",
        "using", "job", "role", "work", "years", "year", "team", "skills",
        "experience", "responsibilities", "responsibility", "about", "ensure",
        "including", "such", "other", "more", "our", "its", "you'll", "we"
    }

    raw = re.findall(r"[a-zA-Z][a-zA-Z+#.-]{1,}", safe_lower(text))
    tokens = [t for t in raw if t not in stopwords and len(t) > 2]
    return tokens


def job_match(resume_text: str, job_text: str, detected_skills):
    if not job_text.strip():
        return None

    resume_lower = safe_lower(resume_text)
    job_lower = safe_lower(job_text)

    matched_skills = []
    missing_skills = []

    for skill, keywords in SKILLS_DATABASE.items():
        mentioned_in_job = any(keyword_in_text(job_lower, k) for k in keywords)
        if mentioned_in_job:
            present_in_resume = any(keyword_in_text(resume_lower, k) for k in keywords)
            if present_in_resume:
                matched_skills.append(skill)
            else:
                missing_skills.append(skill)

    job_tokens = set(tokenize_keywords(job_text))
    resume_tokens = set(tokenize_keywords(resume_text))

    if job_tokens:
        token_overlap = len(job_tokens.intersection(resume_tokens)) / len(job_tokens)
    else:
        token_overlap = 0.0

    skill_total = len(matched_skills) + len(missing_skills)
    skill_overlap = (
        len(matched_skills) / skill_total
        if skill_total else 0.0
    )

    match_score = round((skill_overlap * 70) + (token_overlap * 30))
    match_score = max(0, min(match_score, 100))

    return {
        "score": match_score,
        "matched_skills": matched_skills,
        "missing_skills": missing_skills,
        "keyword_overlap": round(token_overlap * 100),
    }


def generate_basic_feedback(text, word_count, detected_skills, sections):
    feedback = []

    if word_count < 180:
        feedback.append("The resume is quite short. Add relevant projects, achievements, certifications, or stronger experience details.")
    elif word_count <= 700:
        feedback.append("Your resume contains a reasonable amount of content.")
    else:
        feedback.append("The resume is content-heavy. Remove low-value text and keep the document focused.")

    if len(detected_skills) == 0:
        feedback.append("No recognized skills were detected. Add a clear Skills section using the technologies or capabilities you genuinely know.")
    elif len(detected_skills) < 5:
        feedback.append("Only a few recognized skills were detected. Add more relevant skills that you genuinely know.")
    else:
        feedback.append("Good! Several relevant skills were detected.")

    if "Projects" not in sections:
        feedback.append("Consider adding 2–3 relevant projects or portfolio items when appropriate.")

    text_lower = safe_lower(text)
    action_words = [
        "built", "developed", "implemented", "designed", "analyzed",
        "improved", "managed", "created", "optimized", "delivered",
        "reduced", "increased", "automated", "led"
    ]
    action_count = sum(text_lower.count(word) for word in action_words)
    if action_count < 3:
        feedback.append("Use strong action verbs such as Built, Developed, Implemented, Designed, Analyzed, and Improved.")

    return feedback


def detailed_recommendations(text, detected_skills, sections, job_text):
    text_lower = safe_lower(text)
    recs = []

    if "summary / profile" not in {s.lower() for s in sections}:
        recs.append("Add a concise Professional Summary targeted to the type of role you want.")

    if "skills" not in {s.lower() for s in sections}:
        recs.append("Add a dedicated Skills section with clear categories.")

    if "experience" in {s.lower() for s in sections}:
        if not re.search(r"\b\d+(?:\.\d+)?\s?%", text):
            recs.append("Add measurable outcomes to experience bullets where possible (percentages, volume, time saved, revenue, accuracy, scale).")

    if "linkedin" not in text_lower:
        recs.append("Add your LinkedIn profile when relevant to your field.")

    if any(s in {"Python", "C", "C++", "Java", "JavaScript", "SQL", "Machine Learning", "Data Analysis"} for s in detected_skills):
        if "github" not in text_lower:
            recs.append("If you have coding projects, add a GitHub profile or selected project links.")

    if job_text.strip():
        recs.append("Use the target job description to prioritize the skills and keywords that genuinely match your experience.")

    if not recs:
        recs.append("Your resume already has a solid base. Focus on tailoring keywords and impact statements to each target role.")

    return recs


def build_gemini_prompt(resume_text, job_text):
    target = (
        f"""
TARGET JOB DESCRIPTION:
{job_text}
"""
        if job_text.strip()
        else "No target job description was provided. Evaluate the resume broadly and career-agnostically."
    )

    return f"""
You are a professional resume reviewer and ATS optimization specialist.

Review the resume below. Be career-agnostic: do not assume the candidate is applying only to AI/ML, software, logistics, business, or any other single field.

{target}

RESUME:
----------------
{resume_text}
----------------

Return a practical review with exactly these sections:

## 1. Overall Resume Review
Give a concise professional assessment.

## 2. Strengths
List the strongest evidence from the resume.

## 3. Weaknesses
Point out genuine issues, ambiguity, repetition, poor phrasing, missing context, or formatting concerns visible from the extracted text.

## 4. Missing or Recommended Skills
Recommend only skills that are reasonably relevant to the resume or target job. Never tell the candidate to claim a skill they do not actually know.

## 5. ATS Optimization Suggestions
Focus on standard section names, keyword coverage, readability, structure, and parser-friendly formatting.

## 6. Experience & Achievement Improvements
Show how weak bullets could be rewritten using stronger action verbs and measurable impact. Do not invent facts.

## 7. Projects / Portfolio Suggestions
Suggest projects only when relevant to the candidate's field or target job.

## 8. Education & Certification Suggestions
Recommend clearer formatting and placement.

## 9. Keywords to Add
List relevant keywords that are missing or underrepresented.

## 10. Final Recommendations
Give the 3-5 highest-priority actions.

Rules:
- Do not invent work experience, education, skills, dates, employers, metrics, or achievements.
- Clearly say when something is missing.
- If the target job is absent, do not pretend a job match exists.
- Keep the review concise enough to scan.
"""


def generate_gemini_review(resume_text, job_text):
    if not GEMINI_API_KEY:
        return "Gemini is not configured. Add GEMINI_API_KEY to your .env file."

    if genai is None:
        return "The current Google GenAI SDK is not installed. Run: pip install -U google-genai"

    try:
        client = genai.Client(api_key=GEMINI_API_KEY)
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=build_gemini_prompt(resume_text, job_text),
        )
        return response.text or "Gemini returned an empty response."
    except Exception as exc:
        return f"Gemini analysis failed: {exc}"


def make_export_data(state):
    return {
        "app": "AI Resume Analyzer",
        "model": GEMINI_MODEL,
        "resume_name": state.get("resume_name"),
        "word_count": state.get("word_count"),
        "character_count": state.get("character_count"),
        "line_count": state.get("line_count"),
        "detected_skills": state.get("detected_skills"),
        "skill_areas": state.get("skill_areas"),
        "sections": state.get("sections"),
        "ats_score": state.get("ats_score"),
        "ats_label": state.get("ats_label"),
        "job_match": state.get("job_match"),
        "basic_feedback": state.get("basic_feedback"),
        "recommendations": state.get("recommendations"),
    }


# ------------------------------------------------------------
# Sidebar
# ------------------------------------------------------------
with st.sidebar:
    st.markdown("## 📄 AI Resume Analyzer")
    st.caption("Career-agnostic, ATS-focused resume review for students, job seekers, and professionals.")

    st.markdown(
        """
        <div class="feature-card">
            <div style="font-size:.78rem;color:#8fa3bb;text-transform:uppercase;letter-spacing:.08em;font-weight:800;">
                Dashboard
            </div>
            <div style="font-size:1rem;font-weight:800;margin-top:.25rem;">Resume Intelligence</div>
            <div class="tiny" style="margin-top:.35rem;">
                Parse • Score • Match • Improve
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("### Analysis includes")
    checklist = [
        "Skill detection",
        "Resume section detection",
        "Estimated ATS score",
        "Job-description matching",
        "Gemini-powered review",
        "Actionable improvements",
    ]

    for item in checklist:
        st.markdown(f"✅ {item}")

    st.divider()

    with st.expander("⚙️ AI Settings"):
        st.caption(f"Model: `{GEMINI_MODEL}`")
        if GEMINI_API_KEY:
            st.success("Gemini API key detected.")
        else:
            st.warning("Gemini API key not detected.")
        st.caption("The local resume analysis still works without Gemini.")

    st.markdown(
        '<div class="tiny">Privacy: your API key should stay in <code>.env</code> and should never be committed to GitHub.</div>',
        unsafe_allow_html=True,
    )


# ------------------------------------------------------------
# Hero
# ------------------------------------------------------------
st.markdown(
    """
    <div class="aurora-strip"></div>

    <div class="workflow">
        <div class="workflow-step">
            <span class="workflow-num">1</span>
            <div class="workflow-label">Upload</div>
            <div class="tiny">PDF / DOCX</div>
        </div>
        <div class="workflow-step">
            <span class="workflow-num">2</span>
            <div class="workflow-label">Parse</div>
            <div class="tiny">Extract text</div>
        </div>
        <div class="workflow-step">
            <span class="workflow-num">3</span>
            <div class="workflow-label">Analyze</div>
            <div class="tiny">Skills + ATS + fit</div>
        </div>
        <div class="workflow-step">
            <span class="workflow-num">4</span>
            <div class="workflow-label">Improve</div>
            <div class="tiny">Actionable feedback</div>
        </div>
    </div>

    <div class="hero">
        <div class="hero-title">📄 AI Resume Analyzer</div>
        <div class="hero-sub">
            Analyze your resume for skills, structure, ATS readiness, job fit, and practical improvements.
        </div>
        <div class="badge">● Live analysis • Career-agnostic • ATS-focused • AI-powered</div>
    </div>
    """,
    unsafe_allow_html=True,
)


# ------------------------------------------------------------
# Input area
# ------------------------------------------------------------
left, right = st.columns([1.05, 0.95], gap="large")

with left:
    st.markdown('<div class="section-title">🎯 Target Job <span class="tiny">(Optional)</span></div>', unsafe_allow_html=True)
    st.caption("Paste a job description to compare the resume with a specific role.")
    job_text = st.text_area(
        "Target job description",
        height=190,
        placeholder=(
            "Example:\n"
            "We are looking for a Python developer with SQL, Git, REST APIs, "
            "problem-solving and teamwork skills..."
        ),
        label_visibility="collapsed",
    )

with right:
    st.markdown('<div class="section-title">📤 Upload Resume</div>', unsafe_allow_html=True)
    st.caption("PDF or DOCX files work best. Text-based resumes are easiest to parse.")
    uploaded_file = st.file_uploader(
        "Upload your resume",
        type=["pdf", "docx"],
        label_visibility="collapsed",
    )


# ------------------------------------------------------------
# Analyze once file is uploaded
# ------------------------------------------------------------
if uploaded_file is not None:
    st.markdown(
        f"""
        <div class="quick-action">
            <span class="pulse-orb"></span>
            <b>Ready to scan {uploaded_file.name}</b>
            <div class="tiny" style="margin-top:.25rem;">
                {uploaded_file.size / 1024:.1f} KB • Text extraction • Skill detection • ATS estimation
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    ac1, ac2, ac3 = st.columns([1, 1.4, 1])
    with ac2:
        analyze_clicked = st.button(
            "🚀 Analyze My Resume",
            use_container_width=True,
            type="primary",
            key="analyze_resume_cta",
        )

    if analyze_clicked:
        st.session_state["analysis_started"] = True
    elif not st.session_state.get("analysis_started", False):
        st.info("Your file is ready. Click **Analyze My Resume** to start the full scan.")
        st.stop()

if uploaded_file is None:
    st.markdown(
        """
        <div class="empty-state">
            <div style="font-size:2.3rem;">📄</div>
            <div style="font-size:1.15rem;font-weight:800;margin-top:.4rem;">Your analysis starts here</div>
            <div class="tiny" style="margin-top:.35rem;">
                Upload a text-based PDF or DOCX resume to see ATS scoring, skills, sections,
                job matching, and improvement recommendations.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.stop()

resume_text = extract_text(uploaded_file)

if not resume_text:
    st.error("I couldn't extract selectable text from this file. Try a text-based PDF or DOCX.")
    st.stop()

detected_skills = detect_skills(resume_text)
skill_areas = detect_skill_areas(detected_skills)
sections = detect_sections(resume_text)
word_count, character_count, line_count = get_statistics(resume_text)
ats_score = calculate_ats_score(resume_text, detected_skills, sections)
ats_text, ats_kind = ats_label(ats_score)
job_match_result = job_match(resume_text, job_text, detected_skills)
basic_feedback = generate_basic_feedback(
    resume_text,
    word_count,
    detected_skills,
    sections,
)
recommendations = detailed_recommendations(
    resume_text,
    detected_skills,
    sections,
    job_text,
)

state = {
    "resume_name": uploaded_file.name,
    "word_count": word_count,
    "character_count": character_count,
    "line_count": line_count,
    "detected_skills": detected_skills,
    "skill_areas": [{"area": a, "count": c} for a, c in skill_areas],
    "sections": sections,
    "ats_score": ats_score,
    "ats_label": ats_text,
    "job_match": job_match_result,
    "basic_feedback": basic_feedback,
    "recommendations": recommendations,
}

st.success(f"Resume loaded successfully • {uploaded_file.name}")

st.markdown(
    """
    <div class="floating-hint">
        <span class="pulse-orb"></span>
        <b>Live scan complete.</b> The dashboard below is driven by your extracted resume text.
    </div>
    """,
    unsafe_allow_html=True,
)

st.caption("Local analysis is deterministic. ATS scoring is an estimate, and AI suggestions should be verified before adding anything to your resume.")

# ------------------------------------------------------------
# Snapshot
# ------------------------------------------------------------
st.markdown("## 📊 Resume Snapshot")

m1, m2, m3, m4, m5 = st.columns(5)

with m1:
    st.markdown(
        f"""
        <div class="metric-card">
            <div class="metric-label">ATS Score</div>
            <div class="metric-value">{ats_score}/100</div>
            <div class="metric-note">{ats_text}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

with m2:
    st.markdown(
        f"""
        <div class="metric-card">
            <div class="metric-label">Skills</div>
            <div class="metric-value">{len(detected_skills)}</div>
            <div class="metric-note">Recognized</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

with m3:
    st.markdown(
        f"""
        <div class="metric-card">
            <div class="metric-label">Sections</div>
            <div class="metric-value">{len(sections)}</div>
            <div class="metric-note">Detected</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

with m4:
    st.markdown(
        f"""
        <div class="metric-card">
            <div class="metric-label">Words</div>
            <div class="metric-value">{word_count}</div>
            <div class="metric-note">Resume length</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

with m5:
    job_display = f"{job_match_result['score']}%" if job_match_result else "—"
    note = "Target-role match" if job_match_result else "Optional"
    st.markdown(
        f"""
        <div class="metric-card">
            <div class="metric-label">Job Match</div>
            <div class="metric-value">{job_display}</div>
            <div class="metric-note">{note}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


st.markdown("## ✨ Intelligence Overview")

section_score = min(100, round((len(sections) / 8) * 100))
skill_score = min(100, round((len(detected_skills) / 12) * 100))
length_score = 100 if 180 <= word_count <= 700 else 78 if 120 <= word_count <= 950 else 52
lower_resume = safe_lower(resume_text)
impact_score = min(
    100,
    (35 if "experience" in {s.lower() for s in sections} else 0)
    + (35 if re.search(r"\b\d+(?:\.\d+)?\s?%", resume_text) else 0)
    + (30 if re.search(r"\b(?:increased|decreased|reduced|improved|saved|grew|managed|delivered|built|developed)\b", lower_resume) else 0),
)

bento_cols = st.columns(4, gap="medium")
overview_cards = [
    ("ATS Readiness", ats_score, ats_text),
    ("Skill Breadth", skill_score, f"{len(detected_skills)} skills"),
    ("Section Coverage", section_score, f"{len(sections)} sections"),
    ("Impact Signals", impact_score, "Evidence strength"),
]
for col, card in zip(bento_cols, overview_cards):
    label, value, note = card
    with col:
        st.markdown(
            f"""
            <div class="bento-card">
                <div class="bento-label">{label}</div>
                <div class="bento-value">{value}%</div>
                <div class="tiny" style="margin-top:.5rem;">{note}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

st.markdown("<div style='height:.8rem'></div>", unsafe_allow_html=True)

ov1, ov2 = st.columns([.9, 1.1], gap="large")

with ov1:
    st.markdown(
        f"""
        <div class="donut-wrap">
            <div class="donut" style="--pct:{ats_score};">
                <div class="donut-content">
                    <div class="donut-number">{ats_score}</div>
                    <div class="donut-caption">ATS estimate</div>
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

with ov2:
    st.markdown(
        """
        <div class="radar-card">
            <div style="font-size:1rem;font-weight:850;">Resume Signal Radar</div>
            <div class="tiny" style="margin-top:.25rem;">
                Visual summary of the signals used by the analyzer.
            </div>
            <div class="radar-grid">
        """,
        unsafe_allow_html=True,
    )

    radar_values = [
        ("ATS readiness", ats_score),
        ("Skills", skill_score),
        ("Structure", section_score),
        ("Impact", impact_score),
        ("Readability", length_score),
    ]

    for name, score_value in radar_values:
        st.markdown(
            f"""
                <div class="radar-row">
                    <div class="radar-name">{name}</div>
                    <div class="radar-track">
                        <div class="radar-fill" style="--fill:{score_value}%;"></div>
                    </div>
                    <div class="radar-score">{score_value}</div>
                </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown("</div></div>", unsafe_allow_html=True)



st.divider()

# ------------------------------------------------------------
# ATS + Job Match
# ------------------------------------------------------------
ats_col, match_col = st.columns([1.0, 1.0], gap="large")

with ats_col:
    st.markdown("## 📋 Estimated ATS Compatibility")

    css_class = {
        "good": "status-good",
        "mid": "status-mid",
        "low": "status-low",
    }[ats_kind]

    ring_degrees = round(ats_score * 3.6, 1)
    st.markdown(
        f"""
        <div class="panel">
            <div style="display:flex;justify-content:space-between;align-items:center;gap:1.2rem;flex-wrap:wrap;">
                <div class="score-ring" style="background:conic-gradient(#38bdf8 {ring_degrees}deg, #1b2b40 0deg);">
                    <div class="score-ring-inner">
                        <div class="score-ring-number">{ats_score}</div>
                        <div class="score-ring-label">/ 100</div>
                    </div>
                </div>
                <div style="flex:1;min-width:210px;">
                    <div class="{css_class}">{ats_text}</div>
                    <div style="font-size:1.15rem;font-weight:800;margin-top:.65rem;">
                        Estimated ATS compatibility
                    </div>
                    <div class="tiny" style="margin-top:.35rem;">
                        A practical educational estimate — not a score from a real employer ATS.
                    </div>
                </div>
            </div>
            <div style="margin-top:1.15rem;">
                <div style="height:9px;background:#1b2b40;border-radius:999px;overflow:hidden;">
                    <div style="width:{ats_score}%;height:100%;background:linear-gradient(90deg,#0ea5e9,#38bdf8);"></div>
                </div>
            </div>
            <div class="tiny" style="margin-top:.75rem;">
                Based on skills, section structure, contact details, links, impact language, and resume length.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.expander("View score breakdown"):
        text_lower = safe_lower(resume_text)

        if len(detected_skills) >= 8:
            skills_points = 25
        elif len(detected_skills) >= 5:
            skills_points = 21
        elif len(detected_skills) >= 3:
            skills_points = 16
        elif len(detected_skills) >= 1:
            skills_points = 9
        else:
            skills_points = 0

        structure_points = min(
            len({"Contact Information", "Education", "Experience", "Skills"}.intersection(set(sections))) * 5,
            20,
        )
        contact_points = (
            (8 if re.search(r"[\w\.-]+@[\w\.-]+\.\w+", resume_text) else 0)
            + (7 if re.search(r"(?:\+?\d[\d\s().-]{7,}\d)", resume_text) else 0)
        )
        link_points = (
            (5 if "linkedin" in text_lower else 0)
            + (5 if "github" in text_lower or "portfolio" in text_lower else 0)
        )
        impact_points = min(
            (7 if "project" in text_lower else 0)
            + (4 if re.search(r"\b\d+(?:\.\d+)?\s?%", resume_text) else 0)
            + (4 if re.search(r"\b(?:increased|decreased|reduced|improved|saved|grew|managed|delivered)\b", text_lower) else 0),
            15,
        )
        readability_points = 15 if 180 <= word_count <= 700 else 10 if 120 <= word_count <= 950 else 5

        breakdowns = [
            ("Skills coverage", skills_points, 25),
            ("Structure", structure_points, 20),
            ("Contact details", contact_points, 15),
            ("Links / online presence", link_points, 10),
            ("Achievements & impact", impact_points, 15),
            ("Readability & length", readability_points, 15),
        ]

        for label, value, total in breakdowns:
            st.write(f"**{label}: {value}/{total}**")
            st.progress(min(value / total, 1.0))

with match_col:
    st.markdown("## 🎯 Job Description Match")

    if job_match_result is None:
        st.markdown(
            """
            <div class="match-box">
                <b>Paste a target job description above to enable job matching.</b>
                <div class="tiny" style="margin-top:.4rem;">
                    The analyzer will compare resume keywords and recognized skills with the target role.
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        score = job_match_result["score"]
        if score >= 80:
            status = "Strong match"
        elif score >= 60:
            status = "Moderate match"
        else:
            status = "Needs tailoring"

        st.markdown(
            f"""
            <div class="match-box">
                <div style="display:flex;justify-content:space-between;align-items:center;gap:1rem;">
                    <div>
                        <div class="match-score">{score}%</div>
                        <div style="font-size:1.05rem;font-weight:750;margin-top:.35rem;">{status}</div>
                    </div>
                    <div class="{'status-good' if score >= 80 else 'status-mid' if score >= 60 else 'status-low'}">
                        {status}
                    </div>
                </div>
                <div style="height:8px;background:#1b2b40;border-radius:999px;overflow:hidden;margin-top:1rem;">
                    <div style="width:{score}%;height:100%;background:linear-gradient(90deg,#0ea5e9,#38bdf8);"></div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        js1, js2 = st.columns(2)
        with js1:
            st.markdown(
                f'<div class="match-stat"><div class="match-stat-value">{len(job_match_result["matched_skills"])}</div><div class="match-stat-label">Matched skills</div></div>',
                unsafe_allow_html=True,
            )
        with js2:
            st.markdown(
                f'<div class="match-stat"><div class="match-stat-value">{job_match_result["keyword_overlap"]}%</div><div class="match-stat-label">Keyword overlap</div></div>',
                unsafe_allow_html=True,
            )

        if job_match_result["matched_skills"]:
            st.markdown("### ✅ Matched Skills")
            pills = "".join(
                [f'<span class="skill-pill">✓ {skill}</span>' for skill in job_match_result["matched_skills"][:20]]
            )
            st.markdown(pills, unsafe_allow_html=True)

        if job_match_result["missing_skills"]:
            st.markdown("### ⚠️ Missing / Unseen Skills")
            pills = "".join(
                [f'<span class="area-pill">+ {skill}</span>' for skill in job_match_result["missing_skills"][:20]]
            )
            st.markdown(pills, unsafe_allow_html=True)

        st.caption("Only add a missing skill if you genuinely know it.")

st.divider()

# ------------------------------------------------------------
# Skills + Sections
# ------------------------------------------------------------
skills_col, sections_col = st.columns([1.15, 0.85], gap="large")

with skills_col:
    st.markdown("## 🛠️ Skills Detected")
    st.caption(f"Found {len(detected_skills)} recognized skills in your resume.")

    if detected_skills:
        pills = "".join([
            f'<span class="skill-pill" style="animation-delay:{min(i * 0.035, 0.6):.3f}s">✓ {skill}</span>'
            for i, skill in enumerate(detected_skills)
        ])
        st.markdown(pills, unsafe_allow_html=True)
    else:
        st.warning("No recognized skills were detected.")

    if skill_areas:
        st.markdown("### Detected Areas")
        area_pills = "".join(
            [
                f'<span class="area-pill" style="animation-delay:{min(i * 0.06, 0.6):.3f}s">{area} ({count})</span>'
                for i, (area, count) in enumerate(skill_areas)
            ]
        )
        st.markdown(area_pills, unsafe_allow_html=True)

with sections_col:
    st.markdown("## 📚 Resume Sections")
    st.caption(f"Detected {len(sections)} standard sections.")

    if sections:
        section_html = "".join(
            [
                f"""
                <div class="feature-card" style="margin-bottom:.55rem;padding:.75rem .9rem;">
                    <span style="color:#86efac;font-weight:800;">✓</span>
                    <span style="margin-left:.4rem;font-weight:700;">{section}</span>
                </div>
                """
                for section in sections
            ]
        )
        st.markdown(section_html, unsafe_allow_html=True)
    else:
        st.warning("No standard resume sections were detected.")

st.divider()

# ------------------------------------------------------------
# Gemini review
# ------------------------------------------------------------
st.markdown("## 🤖 Gemini Resume Review")
st.caption(
    "The AI review complements the local analysis. It does not replace recruiter judgment or a real ATS."
)

analysis_signature = hashlib.sha256(
    (uploaded_file.name + str(uploaded_file.size) + job_text + resume_text).encode("utf-8", errors="ignore")
).hexdigest()

if st.session_state.get("analysis_signature") != analysis_signature:
    st.session_state.analysis_signature = analysis_signature
    st.session_state.gemini_review = None

if st.button("✨ Generate / Refresh AI Review", use_container_width=True, type="primary"):
    with st.spinner("Gemini is analyzing your resume..."):
        st.session_state.gemini_review = generate_gemini_review(
            resume_text,
            job_text,
        )

if st.session_state.get("gemini_review"):
    with st.container(border=True):
        st.markdown(st.session_state.gemini_review)
else:
    st.info("Generate an AI review after uploading your resume. Local analysis is already available below.")

st.divider()

# ------------------------------------------------------------
# Relevant skill suggestions + Quick feedback
# ------------------------------------------------------------
suggested_skills = [
    skill for skill in COMMON_SUGGESTIONS
    if skill not in detected_skills
]

sugg_col, feedback_col = st.columns([1.0, 1.0], gap="large")

with sugg_col:
    st.markdown("## 💡 Relevant Skill Suggestions")
    st.caption("Potentially useful skills from broadly applicable areas. Only add a skill you genuinely know.")

    if suggested_skills:
        pills = "".join([
                f'<span class="area-pill" style="animation-delay:{min(i * 0.05, 0.6):.3f}s">+ {skill}</span>'
                for i, skill in enumerate(suggested_skills[:12])
            ])
        st.markdown(pills, unsafe_allow_html=True)
    else:
        st.success("Your resume already contains the main suggested skills! 🎉")

with feedback_col:
    st.markdown("## 📝 Quick Feedback")

    for item in basic_feedback:
        st.info(item)

st.divider()

# ------------------------------------------------------------
# Recommendation center
# ------------------------------------------------------------
st.markdown("## 🚀 Improvement Center")

rec1, rec2 = st.columns([1.0, 1.0], gap="large")

with rec1:
    st.markdown("### Priority Recommendations")
    for i, item in enumerate(recommendations[:6], start=1):
        priority_class = "priority-high" if i == 1 else "priority-medium" if i <= 3 else "priority-low"
        priority_label = "HIGH PRIORITY" if i == 1 else "MEDIUM PRIORITY" if i <= 3 else "QUICK WIN"
        st.markdown(
            f"""
            <div class="{priority_class}">
                <div style="font-size:.7rem;font-weight:850;letter-spacing:.08em;color:#9fb0c5;">
                    {priority_label}
                </div>
                <div style="font-weight:700;margin-top:.25rem;">
                    {i}. {item}
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

with rec2:
    st.markdown("### Resume Improvement Checklist")
    checklist_items = [
        "Keep the resume clean and easy to read.",
        "Use a professional email address.",
        "Add LinkedIn when relevant to your field.",
        "Add GitHub / portfolio when relevant to your field.",
        "Include 2–3 strong, relevant projects when appropriate.",
        "Mention relevant skills that you genuinely know.",
        "Use measurable achievements whenever possible.",
        "Use strong action verbs.",
        "Avoid unnecessary personal information.",
        "Keep the resume concise and targeted to the job.",
    ]

    for item in checklist_items:
        st.checkbox(item, value=False, key=f"check_{hash(item)}")

st.divider()

# ------------------------------------------------------------
# Advanced / raw details
# ------------------------------------------------------------
with st.expander("📄 View extracted resume text"):
    st.text_area(
        "Extracted text",
        resume_text,
        height=320,
        label_visibility="collapsed",
    )

with st.expander("📊 View detailed resume statistics"):
    s1, s2, s3 = st.columns(3)
    s1.metric("Words", word_count)
    s2.metric("Characters", character_count)
    s3.metric("Lines", line_count)

# ------------------------------------------------------------
# Export
# ------------------------------------------------------------
st.divider()
st.markdown("## 📦 Export Analysis")

export_json = json.dumps(make_export_data(state), indent=2, ensure_ascii=False)

st.download_button(
    "⬇️ Download analysis JSON",
    data=export_json,
    file_name="resume_analysis_report.json",
    mime="application/json",
    use_container_width=True,
)

st.caption(
    "Privacy note: keep your Gemini API key in .env and never commit the .env file to GitHub."
)

st.markdown(
    """
    <div class="footer">
        Built with Streamlit • Local resume parsing • Gemini AI • Career-agnostic analysis
    </div>
    """,
    unsafe_allow_html=True,
)
